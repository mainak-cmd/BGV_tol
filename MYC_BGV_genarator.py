import os.path
from pathlib import Path
import pandas as pd
import Back_end_logic

import docx

def myc_bgv_doc(key1,key2,key3,key4,key5,key6,key7,key8):
    mydoc = docx.Document()
    mydoc.add_heading("BGV Initiation Request Form", 0)
    table = mydoc.add_table(rows=13, cols=2)

    table.cell(0, 0).text = 'Hire Type - Offered / Existing'
    table.cell(1, 0).text = 'Employee / Candidate Name'
    table.cell(2, 0).text = 'CG Employee ID (Emp ID mandate Existing Employees)'
    table.cell(3, 0).text = 'Candidate ID Candidate ID Mandate for Pre- Hire / Offered'
    table.cell(4, 0).text = 'Employee Capgemini Email ID'
    table.cell(5, 0).text = 'Date of Joining Capgemini'
    table.cell(6, 0).text = 'BGV - Project Name'
    table.cell(7, 0).text = 'Standard BGC Status (Initiated/Notinitiated/Completed)'
    table.cell(8, 0).text = 'Work/Joining Location'
    table.cell(9, 0).text = '**Requestor Email - Emp ID'
    table.cell(10, 0).text = 'Capgemini Entity - (ACIS /FS/ BPO or BSV)'
    table.cell(11, 0).text = 'Purchase Order Number'
    table.cell(12, 0).text = 'Purchase Order Value'

    table.cell(0, 1).text = f"{key1}"
    table.cell(1, 1).text = f"{key2}"
    table.cell(2, 1).text = f"{key3}"
    table.cell(3, 1).text = f"{key4}"
    table.cell(4, 1).text = f"{key5}"
    table.cell(5, 1).text = f"{key6}"
    table.cell(6, 1).text = "HSBC"
    table.cell(8, 1).text = f"{key7}"
    table.cell(10, 1).text = f"{key8}"

    """base_dir = Path(__file__).parent
    path = os.path.join(base_dir, "Office use only")
    if not os.path.exists(path):
        os.makedirs(path)"""
    dir_path = Back_end_logic.Upload_files("Office use only")
    New_path = os.path.join(dir_path,"BGV Initiation Request Form.docx")
    return mydoc.save(f"{New_path}")





